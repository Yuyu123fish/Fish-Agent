"""Word .docx parser based on python-docx."""

from __future__ import annotations

import io

from fish_worker.parser.base import BaseParser, RawElement


class DocxParser(BaseParser):
    """Extract paragraphs, headings, and table rows from .docx documents."""

    def parse(self, content: bytes, filename: str, tmp_dir: str | None = None) -> list[RawElement]:
        from docx import Document

        doc = Document(io.BytesIO(content))
        elements: list[RawElement] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            element_type = "Title" if "heading" in style_name.lower() else "Text"
            elements.append(RawElement(text=text, page=None, element_type=element_type))

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(cells)
                if row_text.strip(" |"):
                    elements.append(RawElement(text=row_text, page=None, element_type="Table"))

        return elements
