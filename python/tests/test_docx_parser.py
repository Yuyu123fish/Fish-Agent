from __future__ import annotations

import io
import unittest

from docx import Document

from fish_worker.parser.docx_parser import DocxParser


def make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("项目标题", level=1)
    doc.add_paragraph("普通段落")
    doc.add_paragraph("")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "角色"
    table.cell(1, 0).text = "小鱼"
    table.cell(1, 1).text = "助手"

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


class DocxParserTest(unittest.TestCase):
    def test_extracts_headings_paragraphs_and_table_rows(self) -> None:
        elements = DocxParser().parse(make_docx_bytes(), "demo.docx")

        self.assertEqual(
            [
                ("项目标题", "Title"),
                ("普通段落", "Text"),
                ("姓名 | 角色", "Table"),
                ("小鱼 | 助手", "Table"),
            ],
            [(e.text, e.element_type) for e in elements],
        )
        self.assertTrue(all(e.page is None for e in elements))


if __name__ == "__main__":
    unittest.main()
